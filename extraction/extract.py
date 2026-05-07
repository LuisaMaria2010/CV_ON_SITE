"""Structured extraction for PDF/DOCX/TXT documents."""
from __future__ import annotations

import io
import logging
import re
from statistics import median

import fitz
from docx import Document

from core.config import settings
from core.errors import TextExtractionError
from domain.document_elements import DocumentElement


logger = logging.getLogger(__name__)


# =========================================================
# Public API
# =========================================================


def extract_elements(file_bytes: bytes, mime_type: str | None = None) -> list[DocumentElement]:
    """Extract structured elements without changing the plain-text API."""
    if not file_bytes:
        return []

    try:
        if _is_pdf(mime_type, file_bytes):
            return _elements_from_pdf(file_bytes)

        if _is_docx(mime_type):
            return _elements_from_docx(file_bytes)

        return _elements_from_txt(file_bytes)
    except Exception as e:
        logger.exception("Structured extraction failed")
        raise TextExtractionError(str(e)) from e


# =========================================================
# Format detection
# =========================================================

def _is_pdf(mime: str | None, data: bytes) -> bool:
    """
    Rileva se il file è un PDF tramite MIME o signature.
    """
    if mime and "pdf" in mime:
        return True
    return data.startswith(b"%PDF")


def _is_docx(mime: str | None) -> bool:
    """
    Rileva se il file è un DOCX tramite MIME type.
    """
    """
    DOCX detection solo via MIME (più sicuro).
    """
    if not mime:
        return False

    return "officedocument" in mime or "word" in mime


# =========================================================
# Extractors
# =========================================================


def _elements_from_pdf(data: bytes) -> list[DocumentElement]:
    font_sizes: list[float] = []
    table_elements: list[DocumentElement] = []
    per_line_items: list[dict] = []

    with fitz.open(stream=data, filetype="pdf") as doc:
        for page_number, page in enumerate(doc, start=1):

            # ── Step 1: native table detection via find_tables() ──────────────
            table_bboxes: list[tuple] = []
            page_dict: dict = {}
            try:
                # get_text("dict") forces creation of the internal TextPage object
                # that find_tables() needs.  Without this, PyMuPDF emits a C-level
                # "not a textpage of this page" message to stderr on some PDFs.
                page_dict = page.get_text("dict")
                tabs = page.find_tables()
                for tab in tabs.tables:
                    bx0, by0, bx1, by1 = tab.bbox
                    table_bboxes.append((bx0, by0, bx1, by1))
                    rows: list[list[str]] = []
                    for row in tab.extract():
                        clean_row = [(cell or "").strip() for cell in row]
                        if any(clean_row):
                            rows.append(clean_row)
                    if len(rows) >= 2 and len(rows[0]) >= 2:
                        table_elements.append(
                            DocumentElement(
                                element_type="table",
                                rows=rows,
                                page_number=page_number,
                                vertical_position=float(by0),
                                horizontal_position=float(bx0),
                            )
                        )
            except Exception:
                logger.debug("find_tables failed on page %d", page_number)
                if not page_dict:
                    page_dict = page.get_text("dict")

            # ── Step 2: collect per-line items with individual line bbox ──────
            # Using per-line y-coordinates allows column-aware reading order:
            # lines from left and right columns that share the same y-level
            # are later merged into a single "skill  level" row.
            # page_dict was already fetched in Step 1; reuse it here.
            for block in page_dict.get("blocks", []):
                if block.get("type") != 0:
                    continue

                bx0, by0, bx1, by1 = block.get("bbox", [0, 0, 0, 0])

                # Skip blocks covered by a detected table
                skip = False
                for tx0, ty0, tx1, ty1 in table_bboxes:
                    ox = max(0.0, min(bx1, tx1) - max(bx0, tx0))
                    oy = max(0.0, min(by1, ty1) - max(by0, ty0))
                    if ox * oy / max(1.0, (bx1 - bx0) * (by1 - by0)) > 0.5:
                        skip = True
                        break
                if skip:
                    continue

                for line in block.get("lines", []):
                    # (x0, x1, text) — track x-coords to detect adjacent spans
                    span_parts: list[tuple[float, float, str]] = []
                    span_sizes: list[float] = []
                    line_bbox = line.get("bbox", [bx0, by0, bx1, by1])
                    for span in line.get("spans", []):
                        raw = span.get("text") or ""
                        if not raw.strip():
                            continue
                        sbbox = span.get("bbox", [0, 0, 0, 0])
                        span_parts.append((float(sbbox[0]), float(sbbox[2]), raw))
                        size = float(span.get("size") or 0)
                        if size > 0:
                            span_sizes.append(size)
                            font_sizes.append(size)

                    if span_parts:
                        # Join consecutive spans: if x-gap < SPAN_GAP_PT the spans
                        # are visually adjacent (same word split across style runs,
                        # e.g. a coloured first letter) → merge WITHOUT space.
                        # A wider gap means an actual inter-word space → keep space.
                        SPAN_GAP_PT = 1.5
                        joined = span_parts[0][2].rstrip()
                        for i in range(1, len(span_parts)):
                            gap = span_parts[i][0] - span_parts[i - 1][1]
                            sep = "" if gap < SPAN_GAP_PT else " "
                            joined += sep + span_parts[i][2].lstrip()
                        line_text = joined.strip()
                        if line_text:
                            per_line_items.append(
                                {
                                    "page_number": page_number,
                                    "y0": float(line_bbox[1]),
                                    "x0": float(line_bbox[0]),
                                    "text": line_text,
                                    "size": sum(span_sizes) / len(span_sizes) if span_sizes else 0.0,
                                }
                            )

    base_size = median(font_sizes) if font_sizes else 11.0

    # ── Step 3: sort per-line items by (page, y, x) ───────────────────────────
    per_line_items.sort(key=lambda l: (l["page_number"], l["y0"], l["x0"]))

    # ── Step 4: group same-y lines into visual rows and merge column cells ────
    # Lines within Y_TOL pts on the same page are considered the same visual row.
    # Multiple cells in a row (from different x-columns) are joined with "  " so
    # that _maybe_parse_table can later detect the table structure.
    Y_TOL = 3.0
    merged_lines: list[dict] = []
    if per_line_items:
        current_row: list[dict] = [per_line_items[0]]
        for item in per_line_items[1:]:
            prev = current_row[0]
            if item["page_number"] == prev["page_number"] and abs(item["y0"] - prev["y0"]) <= Y_TOL:
                current_row.append(item)
            else:
                merged_lines.append(_merge_row_cells(current_row))
                current_row = [item]
        merged_lines.append(_merge_row_cells(current_row))

    # ── Step 5: re-group merged lines into blocks ─────────────────────────────
    # Consecutive merged lines that are close in y (≤ 2.5× size gap) and share the
    # same column structure belong to the same block.
    raw_blocks: list[dict] = []
    if merged_lines:
        current_block: list[dict] = [merged_lines[0]]
        for ml in merged_lines[1:]:
            prev = current_block[-1]
            y_gap = ml["y0"] - prev["y0"]
            max_size = max(ml["size"], prev["size"], 1.0)
            same_page = ml["page_number"] == prev["page_number"]
            prev_two_col = "  " in prev["text"]
            curr_two_col = "  " in ml["text"]
            # Split block if: different page, large Y-gap, column structure change,
            # or a significant font-size jump (heading-sized line isolated from body)
            heading_threshold = base_size * 1.15
            size_jump = (
                (ml["size"] >= heading_threshold) != (prev["size"] >= heading_threshold)
            )
            split = (
                not same_page
                or y_gap > max_size * 2.5
                or prev_two_col != curr_two_col
                or size_jump
            )
            if split:
                raw_blocks.append(_finalize_raw_block(current_block))
                current_block = [ml]
            else:
                current_block.append(ml)
        raw_blocks.append(_finalize_raw_block(current_block))

    # ── Step 6: convert raw_blocks → DocumentElements ────────────────────────
    elements: list[DocumentElement] = list(table_elements)

    for block in raw_blocks:
        lines = [line["text"] for line in block["lines"] if line["text"]]
        if not lines:
            continue

        table_rows = _maybe_parse_table(lines)
        if table_rows:
            elements.append(
                DocumentElement(
                    element_type="table",
                    rows=table_rows,
                    page_number=block["page_number"],
                    vertical_position=block["vertical_position"],
                    horizontal_position=block["horizontal_position"],
                )
            )
            continue

        avg_size = sum(line["size"] for line in block["lines"] if line["size"] > 0) / max(
            1,
            sum(1 for line in block["lines"] if line["size"] > 0),
        )

        if len(lines) == 1:
            line = lines[0]
            list_info = _extract_list_info(line)
            if list_info:
                elements.append(
                    DocumentElement(
                        element_type="list_item",
                        text=list_info["body"],
                        level=list_info["level"],
                        list_ordered=list_info["ordered"],
                        page_number=block["page_number"],
                        vertical_position=block["vertical_position"],
                        horizontal_position=block["horizontal_position"],
                    )
                )
                continue

            if _is_heading_candidate(line, avg_size, base_size):
                elements.append(
                    DocumentElement(
                        element_type="heading",
                        text=_clean_text(line),
                        level=_heading_level(line, avg_size, base_size),
                        page_number=block["page_number"],
                        vertical_position=block["vertical_position"],
                        horizontal_position=block["horizontal_position"],
                    )
                )
                continue

        paragraph = _clean_text("\n".join(lines))
        if paragraph:
            elements.append(
                DocumentElement(
                    element_type="paragraph",
                    text=paragraph,
                    page_number=block["page_number"],
                    vertical_position=block["vertical_position"],
                    horizontal_position=block["horizontal_position"],
                )
            )

    # Elements are already in reading order (sorted by y during per-line step)
    return _coalesce_paragraphs(elements)


def _merge_row_cells(row: list[dict]) -> dict:
    """Merge multiple per-line items that share the same y-level into one merged line.

    Items from different x-columns are joined with '  ' (double space) so that
    _maybe_parse_table can later split them into table cells.
    """
    if len(row) == 1:
        return row[0]
    sorted_cells = sorted(row, key=lambda c: c["x0"])
    parts = [c["text"] for c in sorted_cells if c["text"]]
    merged_text = "  ".join(parts)
    avg_size = sum(c["size"] for c in row) / max(1, len(row))
    return {
        "page_number": row[0]["page_number"],
        "y0": row[0]["y0"],
        "x0": row[0]["x0"],
        "text": merged_text,
        "size": avg_size,
    }


def _finalize_raw_block(lines: list[dict]) -> dict:
    """Convert a list of merged line items into a raw_block dict."""
    return {
        "page_number": lines[0]["page_number"],
        "vertical_position": lines[0]["y0"],
        "horizontal_position": lines[0]["x0"],
        "lines": [{"text": l["text"], "size": l["size"]} for l in lines],
    }


def _elements_from_docx(data: bytes) -> list[DocumentElement]:
    file_like = io.BytesIO(data)
    document = Document(file_like)
    elements: list[DocumentElement] = []

    for paragraph in document.paragraphs:
        text = _clean_text(paragraph.text)
        if not text:
            continue

        style_name = (getattr(paragraph.style, "name", "") or "").lower()
        if style_name.startswith("heading"):
            match = re.search(r"(\d+)$", style_name)
            level = int(match.group(1)) if match else 1
            elements.append(DocumentElement(element_type="heading", text=text, level=min(max(level, 1), 6)))
            continue

        list_info = _extract_list_info(text)
        if list_info or "list" in style_name:
            info = list_info or {"body": text, "ordered": False, "level": 1}
            elements.append(
                DocumentElement(
                    element_type="list_item",
                    text=info["body"],
                    level=info["level"],
                    list_ordered=info["ordered"],
                )
            )
            continue

        elements.append(DocumentElement(element_type="paragraph", text=text))

    return _coalesce_paragraphs(elements)


def _elements_from_txt(data: bytes) -> list[DocumentElement]:
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("latin-1", errors="ignore")
    sections = [chunk.strip() for chunk in re.split(r"\n\s*\n", text) if chunk.strip()]
    elements: list[DocumentElement] = []

    for section in sections:
        if "\n" not in section and _is_heading_candidate(section, 14.0, 11.0):
            elements.append(DocumentElement(element_type="heading", text=_clean_text(section), level=_heading_level(section, 14.0, 11.0)))
            continue

        list_info = _extract_list_info(section)
        if list_info and "\n" not in section:
            elements.append(
                DocumentElement(
                    element_type="list_item",
                    text=list_info["body"],
                    level=list_info["level"],
                    list_ordered=list_info["ordered"],
                )
            )
            continue

        elements.append(DocumentElement(element_type="paragraph", text=_clean_text(section)))

    return _coalesce_paragraphs(elements)


# =========================================================
# Cleanup (CRUCIALE per LLM)
# =========================================================

_whitespace_re = re.compile(r"[ \t]+")
_multiline_re = re.compile(r"\n{3,}")
# Matches abbreviations like "S . R . L" where single uppercase letters are
# separated by spaces and dots (e.g. Italian "S.r.l.").  The negative lookahead
# avoids merging initials followed by a full word ("V . Monti").
_spaced_abbrev_re = re.compile(r"([A-Z]) \. ([A-Z])(?![a-zA-Z])")
_unordered_list_re = re.compile(r"^\s*(?P<marker>[•\-\*])\s+(?P<body>.+)$")
_ordered_list_re = re.compile(r"^\s*(?P<index>\d+)[\.)]\s+(?P<body>.+)$")
_heading_numbering_re = re.compile(r"^\d+(?:\.\d+){0,4}\s+")


def _fix_spaced_abbreviations(text: str) -> str:
    """Collapse spurious spaces in abbreviations: 'S . R . L' → 'S.R.L'.

    Applies the pattern repeatedly until stable (handles chains like S . R . L . I).
    Only collapses when the second letter is NOT followed by more letters, so
    initials before full words ('V . Monti') are left unchanged.
    """
    prev = None
    while prev != text:
        prev = text
        text = _spaced_abbrev_re.sub(r"\1.\2", text)
    return text


def _clean_text(text: str) -> str:
    """
    Pulisce e normalizza il testo estratto:
    - rimuove null chars
    - normalizza spazi e newline
    - tronca a max_text_chars
    """
    if not text:
        return ""

    # rimuove null chars (comuni nei PDF)
    text = text.replace("\x00", "")

    # abbreviazioni con spazi attorno ai punti: "S . R . L" → "S.R.L"
    text = _fix_spaced_abbreviations(text)

    # spazi multipli
    text = _whitespace_re.sub(" ", text)

    # newline eccessivi
    text = _multiline_re.sub("\n\n", text)

    text = text.strip()

    # limit token safety
    max_chars = settings.max_text_chars

    if len(text) > max_chars:
        logger.warning(
            "Text truncated to %s chars for token safety",
            max_chars,
        )
        text = text[:max_chars]

    return text


def _extract_list_info(text: str) -> dict | None:
    unordered_match = _unordered_list_re.match(text)
    if unordered_match:
        return {
            "body": _clean_text(unordered_match.group("body")),
            "ordered": False,
            "level": 1,
        }

    ordered_match = _ordered_list_re.match(text)
    if ordered_match:
        return {
            "body": _clean_text(ordered_match.group("body")),
            "ordered": True,
            "level": 1,
        }

    return None


def _is_heading_candidate(text: str, font_size: float, base_size: float) -> bool:
    compact = _clean_text(text)
    if not compact:
        return False
    if len(compact) > 120:
        return False
    if compact.endswith((".", ";", ",")):
        return False
    if _extract_list_info(compact):
        return False
    if _heading_numbering_re.match(compact):
        return True
    if font_size >= base_size * 1.2:
        return True
    words = compact.split()
    if len(words) <= 10 and compact.isupper():
        return True
    return False


def _heading_level(text: str, font_size: float, base_size: float) -> int:
    compact = _clean_text(text)
    numbering_match = re.match(r"^(\d+(?:\.\d+)*)", compact)
    if numbering_match:
        return min(6, max(1, len(numbering_match.group(1).split("."))))
    ratio = font_size / max(base_size, 1.0)
    if ratio >= 1.8:
        return 1
    if ratio >= 1.5:
        return 2
    return 3


def _maybe_parse_table(lines: list[str]) -> list[list[str]] | None:
    if len(lines) < 2:
        return None

    rows: list[list[str]] = []
    widths: set[int] = set()
    for line in lines:
        cells = [cell.strip() for cell in re.split(r"\s{2,}|\t+", line) if cell.strip()]
        if len(cells) < 2:
            return None
        widths.add(len(cells))
        rows.append(cells)

    if len(widths) > 2:
        return None

    target_width = max(widths)
    return [row + [""] * (target_width - len(row)) for row in rows]


def _coalesce_paragraphs(elements: list[DocumentElement]) -> list[DocumentElement]:
    merged: list[DocumentElement] = []
    for element in elements:
        if (
            merged
            and element.element_type == "paragraph"
            and merged[-1].element_type == "paragraph"
            and merged[-1].page_number == element.page_number
        ):
            merged[-1].text = _clean_text(f"{merged[-1].text}\n\n{element.text}")
            continue
        merged.append(element)
    return merged
