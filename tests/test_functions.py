import requests
import json
import os
from docx import Document

BASE = "http://127.0.0.1:7071"

# create a simple CV docx
cv_path = os.path.join("tests", "sample_cv.docx")
if not os.path.exists("tests"):
    os.makedirs("tests")

d = Document()
d.add_heading('Mario Rossi', level=1)
d.add_paragraph('Email: mario.rossi@example.com')
d.add_paragraph('Phone: +39 333 1234567')
d.add_paragraph('Profile: Senior Software Engineer with 10+ years experience in Python and Azure.')
d.add_heading('Experience', level=2)
d.add_paragraph('2020-2026 - Senior Engineer at ACME Corp')
d.add_heading('Education', level=2)
d.add_paragraph('MSc Computer Science, University X')
d.save(cv_path)

print('Created sample CV:', cv_path)

# Test /api/extract
files = {'file': open(cv_path, 'rb')}
try:
    r = requests.post(f"{BASE}/api/extract", files=files, timeout=30)
    print('/api/extract status', r.status_code)
    try:
        print('extract response:', json.dumps(r.json(), indent=2))
    except Exception:
        print('non-json response:', r.text)
except Exception as e:
    print('extract request failed:', e)

# Test /api/backfill/incoming-cv (dry_run)
try:
    r2 = requests.post(f"{BASE}/api/backfill/incoming-cv", json={"dry_run": True, "max_items": 10}, timeout=30)
    print('/api/backfill status', r2.status_code)
    try:
        print('backfill response:', json.dumps(r2.json(), indent=2))
    except Exception:
        print('non-json response:', r2.text)
except Exception as e:
    print('backfill request failed:', e)
